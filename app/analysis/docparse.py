"""Extract plain text from uploaded resumes/letters: PDF, DOCX, TXT."""
import asyncio
import io

MAX_CHARS = 30000


def _parse_pdf(data: bytes) -> str:
    import pdfplumber

    out: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages[:20]:
            out.append(page.extract_text() or "")
    return "\n".join(out)


def _parse_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _parse_sync(data: bytes, filename: str) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        text = _parse_pdf(data)
    elif name.endswith(".docx"):
        text = _parse_docx(data)
    elif name.endswith((".txt", ".md")):
        text = data.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"unsupported file type: {filename}")
    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not text.strip():
        raise ValueError("no text extracted")
    return text[:MAX_CHARS]


async def parse_document(data: bytes, filename: str) -> str:
    """Raises ValueError for unsupported/empty documents."""
    return await asyncio.to_thread(_parse_sync, data, filename)
